"""Build a clean initial-submission DOCX from the audited Markdown article."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "NMI_ARTICLE_DRAFT.md"
OUTPUT = ROOT / "NMI_ARTICLE_SUBMISSION.docx"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_inline(paragraph, text: str) -> None:
    # Preserve code spans and bold markers without exposing Markdown syntax.
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 6),
        ("Heading 2", 13, "2E74B5", 14, 5),
        ("Heading 3", 12, "1F4D78", 10, 4),
    ]:
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Article subtitle" not in [style.name for style in document.styles]:
        subtitle = document.styles.add_style("Article subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = document.styles["Article subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(10)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(89, 89, 89)
    subtitle.paragraph_format.space_after = Pt(14)


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Initial submission draft")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_styles(document)
    add_footer(document)

    first_title = True
    skip_metadata = False
    in_abstract = False
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if first_title and line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line[2:].strip())
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(11, 37, 69)
            first_title = False
            skip_metadata = True
            continue
        if skip_metadata and (line.startswith("**Working Article Draft") or line.startswith("**Status:") or line.startswith("**Target format:")):
            continue
        if line == "## Abstract":
            paragraph = document.add_paragraph("Abstract", style="Heading 1")
            in_abstract = True
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading == "Introduction":
                # Nature's Article format uses an unheaded introduction.
                in_abstract = False
                continue
            document.add_paragraph(heading, style="Heading 1")
            in_abstract = False
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            continue
        if line.startswith("#### "):
            document.add_paragraph(line[5:].strip(), style="Heading 3")
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, line[2:].strip())
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\. ", "", line))
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style="Intense Quote")
            add_inline(paragraph, line[2:].strip())
            continue
        if line.startswith("```"):
            continue
        paragraph = document.add_paragraph()
        add_inline(paragraph, line)

    document.core_properties.title = "Reusable Internal Computation in Language Models"
    document.core_properties.subject = "Nature Machine Intelligence initial submission draft"
    document.core_properties.author = ""
    document.core_properties.comments = "Generated from audited local research package"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
