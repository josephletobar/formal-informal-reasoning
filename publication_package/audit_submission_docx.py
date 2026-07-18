"""Structural audit for the initial-submission Word manuscript."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
doc = Document(ROOT / "NMI_ARTICLE_SUBMISSION.docx")
text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
words = re.findall(r"\b[\w'-]+\b", text)

assert doc.paragraphs[0].text.startswith("Reusable Internal Computation in Language Models")
assert "Abstract" in text
assert "Results" in text
assert "Discussion" in text
assert "Methods" in text
assert "Data and code availability" in text
assert not re.search(r"^#{1,4} |```|\[.*?\]\(https?://", text, flags=re.MULTILINE)
assert len(words) > 2000, len(words)
assert len(text) < 25000, len(text)

print(f"submission DOCX structural audit passed: paragraphs={len(doc.paragraphs)} words={len(words)}")
